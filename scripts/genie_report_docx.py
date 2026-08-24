#!/usr/bin/env python3
"""Create a publication-style DOCX report from a GENIE Markdown report.

The Markdown report remains the reproducible text artifact; this script creates
an accessible, image-rich Word deliverable and appends the exported GENIE tables.
"""
from __future__ import annotations

import argparse
import csv
import html
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Sequence

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - clear CLI error
    raise SystemExit("python-docx is required; install it with: python -m pip install python-docx") from exc

U_PLUS_RE = re.compile(r"<U\+[0-9A-Fa-f]{4,6}>")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")

CORE_FIGURES = [
    ("nmi_before_after.png", "图 1  NMI 初始值与最终值比较", "柱状图展示各分析层面的 NMI 变化；变化值按百分点解释。"),
    ("item_reduction_by_type.png", "图 2  题项数量削减前后比较", "比较完整候选题池与 GENIE 筛查后保留题项数量。"),
    ("removal_waterfall.png", "图 3  UVA 与 bootEGA 的题项流转", "显示冗余筛查、稳定性筛查和最终保留题项的数量关系。"),
    ("attribute_community_heatmap.png", "图 4  预设属性与 EGA 社区对应关系", "热图用于观察预设属性在网络社区中的聚集与错位。"),
]


def decode_unicode_escapes(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        try:
            return chr(int(match.group(0)[3:-1], 16))
        except ValueError:
            return match.group(0)
    return U_PLUS_RE.sub(repl, text)


def clean(text: object) -> str:
    value = decode_unicode_escapes("" if text is None else str(text))
    return value.replace("\\u", "\\u")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value: object, *, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean(value))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = "B7C9D6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1); run._r.append(instr); run._r.append(fld_char2)
    paragraph.add_run(" 页")


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " TOC \\o \"1-3\" \\h \\z \\u "
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "目录将在 Word 中更新"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in [("Title", 24, "1F4E79"), ("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "2F75B5"), ("Heading 3", 11, "404040")]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"; style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
    # A custom caption style makes figure/table labels consistent.
    if "CaptionCN" not in [s.name for s in styles]:
        cap = styles.add_style("CaptionCN", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Microsoft YaHei"; cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        cap.font.size = Pt(9); cap.font.italic = True; cap.font.color.rgb = RGBColor(89, 89, 89)
        cap.paragraph_format.space_after = Pt(6)
    header = sec.header.paragraphs[0]
    header.text = "GENIE 语义筛查结果报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Microsoft YaHei"; run.font.size = Pt(8); run.font.color.rgb = RGBColor(128, 128, 128)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Microsoft YaHei"; run.font.size = Pt(8)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(clean(re.sub(r"^#+\s*", "", text).strip()), level=level)


def add_rich_paragraph(doc: Document, text: str) -> None:
    text = clean(text.strip())
    if not text: return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    # Keep basic Markdown emphasis without exposing syntax.
    pos = 0
    for m in re.finditer(r"(\*\*|`)(.+?)(?:\1)", text):
        if m.start() > pos: p.add_run(text[pos:m.start()])
        r = p.add_run(m.group(2)); r.bold = m.group(1) == "**"
        if m.group(1) == "`": r.font.name = "Consolas"
        pos = m.end()
    if pos < len(text): p.add_run(text[pos:])


def parse_md_table(lines: Sequence[str], start: int):
    if start + 1 >= len(lines) or "|" not in lines[start] or not TABLE_SEPARATOR_RE.match(lines[start + 1]):
        return None, start
    def split(line):
        s = line.strip().strip("|")
        return [clean(x.strip()) for x in s.split("|")]
    headers = split(lines[start]); rows = []; i = start + 2
    while i < len(lines) and "|" in lines[i] and lines[i].strip():
        rows.append(split(lines[i])); i += 1
    return (headers, rows), i


def add_markdown_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=1, cols=max(1, len(headers)))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers): set_cell_text(table.rows[0].cells[j], h, bold=True, size=7)
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells: set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for j in range(len(headers)): set_cell_text(cells[j], row[j] if j < len(row) else "", size=7)
    set_table_borders(table)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, note: str = "") -> bool:
    if not path.is_file(): return False
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try: p.add_run().add_picture(str(path), width=Cm(15.8))
    except Exception as exc:
        add_rich_paragraph(doc, f"（图表无法嵌入：{path.name}；{exc}）"); return False
    cap = doc.add_paragraph(style="CaptionCN"); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.add_run(clean(caption))
    if note:
        n = doc.add_paragraph(style="CaptionCN"); n.alignment = WD_ALIGN_PARAGRAPH.CENTER; n.add_run(clean(note))
    return True


def resolve_report_asset(report_dir: Path, markdown_target: str) -> Path:
    """Resolve a Markdown image target without assuming the host path separator."""
    target = markdown_target.strip().strip("<>")
    target_path = Path(target)
    if target_path.is_absolute():
        return target_path
    # Markdown uses forward slashes even on Windows. Normalize both forms
    # before joining so reports are portable across operating systems.
    return report_dir.joinpath(*target.replace("\\", "/").split("/"))


def add_csv_appendix(doc: Document, path: Path, title: str, max_rows: int = 250) -> None:
    if not path.is_file(): return
    doc.add_heading(title, level=2)
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f); rows = list(reader)
    if not rows:
        add_rich_paragraph(doc, "（无记录）"); return
    headers, data = rows[0], rows[1:max_rows + 1]
    # Keep very wide exports readable by selecting key columns first.
    if len(headers) > 12:
        preferred = ["ID", "statement", "attribute", "type", "EGA_com", "analysis_level", "analysis_type", "stage"]
        idx = [headers.index(c) for c in preferred if c in headers]
        if idx: headers, data = [headers[i] for i in idx], [[r[i] if i < len(r) else "" for i in idx] for r in data]
    add_markdown_table(doc, headers, data)
    if len(rows) - 1 > max_rows: add_rich_paragraph(doc, f"（表格过长，正文展示前 {max_rows} 行；完整数据见 {path.name}。）")


def build_doc(markdown_path: Path, output_path: Path, report_dir: Path, require_core_figures: bool = False) -> None:
    markdown = markdown_path.read_text(encoding="utf-8-sig")
    if U_PLUS_RE.search(markdown): raise ValueError("Markdown contains unresolved <U+....> escapes")
    doc = Document(); configure_document(doc)
    title = "GENIE 语义筛查与网络心理测量结果报告"
    doc.add_paragraph(title, style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("AIGENIE / local_GENIE · 论文式分析报告\n").bold = True
    p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    p.add_run(f"报告源文件：{markdown_path.name}")
    doc.add_page_break()
    doc.add_heading("目录", level=1); toc = doc.add_paragraph(); add_toc(toc); doc.add_page_break()

    lines = markdown.splitlines(); i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue
        if line.startswith("#"):
            level = min(3, len(line) - len(line.lstrip("#")))
            add_heading(doc, line, level); i += 1; continue
        table, nxt = parse_md_table(lines, i)
        if table:
            add_markdown_table(doc, *table); i = nxt; continue
        image = IMAGE_RE.search(line)
        if image:
            add_image(doc, resolve_report_asset(report_dir, image.group(2)), image.group(1)); i += 1; continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet"); p.add_run(clean(line[2:])); i += 1; continue
        if line.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); r = p.add_run(clean(line[2:])); r.italic = True; i += 1; continue
        # Skip Markdown horizontal rules and link-only figure lists; figures are embedded below.
        if re.match(r"^[-*_]{3,}$", line) or (LINK_RE.fullmatch(line) and "figures/" in line): i += 1; continue
        add_rich_paragraph(doc, line); i += 1

    fig_dir = report_dir / "figures"
    doc.add_page_break(); doc.add_heading("图表汇总", level=1)
    missing_core = [name for name, _, _ in CORE_FIGURES if not (fig_dir / name).is_file()]
    if require_core_figures and missing_core:
        raise ValueError("missing core report figures: " + ", ".join(missing_core))
    failed_core = []
    for name, caption, note in CORE_FIGURES:
        if (fig_dir / name).is_file() and not add_image(doc, fig_dir / name, caption, note):
            failed_core.append(name)
    if require_core_figures and failed_core:
        raise ValueError("could not embed core report figures: " + ", ".join(failed_core))
    for path in sorted(fig_dir.glob("*.png")) if fig_dir.is_dir() else []:
        if path.name in {x[0] for x in CORE_FIGURES}: continue
        caption = "附图  网络结构或稳定性诊断：" + path.stem.replace("_", " ")
        add_image(doc, path, caption, "该图为 AIGENIE 返回的 type-level 或 overall 诊断图，具体层级见文件名。")

    doc.add_page_break(); doc.add_heading("附录：可复现数据表", level=1)
    appendix = [("genie_metrics_summary.csv", "附录 A  核心指标明细"), ("genie_final_items.csv", "附录 B  primary final items"), ("genie_type_level_final_items.csv", "附录 C  type-level final items"), ("genie_overall_final_items.csv", "附录 D  overall final items"), ("genie_removed_items.csv", "附录 E  删除题项"), ("genie_redundant_pairs.csv", "附录 F  UVA 冗余题对"), ("genie_warnings.csv", "附录 G  warnings" )]
    for file_name, title_text in appendix: add_csv_appendix(doc, report_dir / file_name, title_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main(argv: Iterable[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Render a GENIE Markdown report as a publication-style DOCX.")
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--report-dir", default=None)
    parser.add_argument("--require-core-figures", action="store_true", help="Fail unless all standard GENIE figures are present and embeddable")
    args = parser.parse_args(list(argv) if argv is not None else None)
    md = Path(args.markdown).resolve(); out = Path(args.output).resolve(); report_dir = Path(args.report_dir).resolve() if args.report_dir else md.parent
    try: build_doc(md, out, report_dir, require_core_figures=args.require_core_figures)
    except (OSError, ValueError) as exc: print(f"DOCX generation failed: {exc}", file=sys.stderr); return 1
    print(f"Generated {out}"); return 0

if __name__ == "__main__": raise SystemExit(main())
