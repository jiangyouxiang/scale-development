#!/usr/bin/env python3
"""Convert a Markdown dimension-structure report to a simple .docx file.

This helper intentionally depends only on python-docx so it can run in the
bundled Codex Python environment. It supports headings, paragraphs, bullets,
numbered lists, block quotes, fenced code blocks, and basic pipe tables.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except Exception as exc:  # pragma: no cover - user-facing dependency message
    print("ERROR: python-docx is required to create .docx reports.", file=sys.stderr)
    print("Install it with: python -m pip install python-docx", file=sys.stderr)
    print(f"Original import error: {exc}", file=sys.stderr)
    raise SystemExit(2)


def clean_inline(text: str) -> str:
    """Remove common Markdown inline markers while preserving readable text."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i]
        if "|" not in line or not line.strip().startswith("|"):
            break
        if not is_table_separator(line):
            rows.append([clean_inline(c) for c in line.strip().strip("|").split("|")])
        i += 1
    return rows, i


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    if "Block Quote" not in styles:
        style = styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
        style.font.italic = True
        style.font.size = Pt(10)
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    ensure_styles(doc)

    in_code = False
    code_buffer: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="No Spacing")
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer.clear()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.lstrip().startswith("|") and "|" in line:
            rows, next_i = parse_table(lines, i)
            if rows:
                max_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
                i = next_i
                continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = clean_inline(heading.group(2))
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=level)
            i += 1
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if bullet:
            doc.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
            i += 1
            continue

        numbered = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if numbered:
            doc.add_paragraph(clean_inline(numbered.group(1)), style="List Number")
            i += 1
            continue

        quote = re.match(r"^>\s?(.*)$", line)
        if quote:
            doc.add_paragraph(clean_inline(quote.group(1)), style="Block Quote")
            i += 1
            continue

        doc.add_paragraph(clean_inline(line))
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Convert a Markdown report to .docx")
    parser.add_argument("markdown", type=Path, help="Input Markdown file")
    parser.add_argument("docx", type=Path, help="Output .docx file")
    args = parser.parse_args(argv)

    if not args.markdown.exists():
        print(f"ERROR: Markdown file not found: {args.markdown}", file=sys.stderr)
        return 1
    convert(args.markdown, args.docx)
    print(f"Wrote {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
